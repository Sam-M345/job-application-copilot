import io
import sys
from pathlib import Path

from docx import Document

COPILOT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(COPILOT_ROOT))

from src.resume_docx import (  # noqa: E402
    MIN_RESUME_WORDS,
    build_resume_docx,
    count_resume_words,
    resume_download_filename,
)
from src.schemas import (  # noqa: E402
    CustomizedResumeContent,
    ResumeJobEntry,
    ResumeSectionContent,
)

_EXAMPLE_PROFILE = {
    "first_name": "Alex",
    "last_name": "Chen",
    "location": "San Francisco, CA",
    "email": "alex.chen.example@gmail.com",
    "phone": "555-123-4567",
    "linkedin": "https://www.linkedin.com/in/example-profile/",
    "portfolio": "https://example-portfolio.dev/",
    "kaggle": "https://www.kaggle.com/example/code",
    "github": "https://github.com/example-user",
    "certifications": [
        {
            "name": "MIT Machine Learning Certificate",
            "url": "https://courses.edx.org/certificates/example-mit",
        },
        {
            "name": "Google Data Analytics Certificate",
            "url": "https://www.coursera.org/account/accomplishments/professional-cert/example",
        },
    ],
}


def test_build_resume_docx_returns_bytes():
    content = CustomizedResumeContent(
        company_name="Acme AI",
        sections=[
            ResumeSectionContent(
                heading="Summary",
                paragraphs=["Applied AI engineer with RAG and LangGraph experience."],
            ),
            ResumeSectionContent(
                heading="Experience",
                job_entries=[
                    ResumeJobEntry(
                        title_line="Senior Analyst | Acme Corp",
                        date_range="Jan 2020 - Present",
                        location="San Francisco, CA",
                        bullet_points=["Built **customer-facing** AI workflows."],
                    )
                ],
            ),
        ],
    )
    data = build_resume_docx(content, profile=_EXAMPLE_PROFILE)
    assert data[:2] == b"PK"
    assert len(data) > 1000


def test_build_resume_docx_header_links_from_profile():
    content = CustomizedResumeContent(
        company_name="Acme AI",
        sections=[
            ResumeSectionContent(
                heading="Summary",
                paragraphs=["Applied AI engineer with RAG and LangGraph experience."],
            )
        ],
    )
    doc = Document(io.BytesIO(build_resume_docx(content, _EXAMPLE_PROFILE)))
    joined = " ".join(p.text for p in doc.paragraphs)
    assert "Alex Chen" in doc.paragraphs[0].text
    assert "https://www.linkedin.com/in/example-profile/" not in joined
    assert "LinkedIn" in joined
    assert "GitHub" in joined
    assert "Kaggle" in joined
    assert "MIT Machine Learning Certificate" in joined
    assert "Google Data Analytics Certificate" in joined
    assert "hyperlink" in doc.element.xml
    rel_targets = [rel.target_ref for rel in doc.part.rels.values() if "hyperlink" in rel.reltype]
    assert "https://github.com/example-user" in rel_targets
    assert "https://www.kaggle.com/example/code" in rel_targets


def test_count_resume_words():
    content = CustomizedResumeContent(
        company_name="Acme",
        sections=[
            ResumeSectionContent(
                heading="Summary",
                paragraphs=["one two three four five"],
            )
        ],
    )
    assert count_resume_words(content) == 5
    assert MIN_RESUME_WORDS == 1100


def test_resume_docx_applies_formatting_rules():
    content = CustomizedResumeContent(
        company_name="Acme",
        sections=[
            ResumeSectionContent(
                heading="experience",
                job_entries=[
                    ResumeJobEntry(
                        title_line="Data Scientist | Acme",
                        date_range="2020 - 2024",
                        bullet_points=["Improved KPIs by **40%**."],
                    )
                ],
            ),
        ],
    )
    doc = Document(io.BytesIO(build_resume_docx(content, profile=_EXAMPLE_PROFILE)))
    section = doc.sections[0]
    assert section.page_width.inches == 8.5
    assert section.left_margin.inches == 0.75
    assert section.top_margin.inches == 0.625

    texts = [p.text for p in doc.paragraphs]
    assert any("EXPERIENCE" in t for t in texts)
    assert any("Data Scientist | Acme" in t and "2020 - 2024" in t for t in texts)
    assert any("40%" in t for t in texts)

    fonts = {
        run.font.name
        for para in doc.paragraphs
        for run in para.runs
        if run.font.name
    }
    assert fonts == {"Arial"}


def test_resume_download_filename():
    content = CustomizedResumeContent(company_name="Acme Corp!")
    name = resume_download_filename(content, {"first_name": "Alex", "last_name": "Chen"})
    assert name == "Alex-Chen-Resume-Acme-Corp.docx"
