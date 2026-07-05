import io
import sys
from pathlib import Path

from docx import Document

COPILOT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(COPILOT_ROOT))

from src.parsers import ParseError, extract_docx_text, parse_document_text, parse_optional_document_text  # noqa: E402


def _sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Alex Chen")
    doc.add_paragraph("Applied AI engineer with LangGraph experience.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_text():
    text = extract_docx_text(_sample_docx_bytes())
    assert "Alex Chen" in text
    assert "LangGraph" in text


def test_parse_resume_accepts_docx_upload():
    text = parse_document_text(
        pasted="",
        uploaded_bytes=_sample_docx_bytes(),
        uploaded_filename="resume.docx",
        label="Resume",
        allowed_extensions=(".pdf", ".docx"),
    )
    assert "Applied AI" in text


def test_parse_optional_document_returns_empty_when_missing():
    assert parse_optional_document_text(
        pasted="",
        uploaded_bytes=None,
        label="Duty Statement",
    ) == ""


def test_parse_optional_document_accepts_docx_upload():
    text = parse_optional_document_text(
        pasted="",
        uploaded_bytes=_sample_docx_bytes(),
        uploaded_filename="duty.docx",
        label="Duty Statement",
    )
    assert "Alex Chen" in text


def test_parse_rejects_docx_for_jd_when_pdf_only():
    try:
        parse_document_text(
            pasted="",
            uploaded_bytes=_sample_docx_bytes(),
            uploaded_filename="jd.docx",
            label="Job description",
            allowed_extensions=(".pdf",),
        )
        raise AssertionError("expected ParseError")
    except ParseError as exc:
        assert "pdf" in str(exc).lower()
