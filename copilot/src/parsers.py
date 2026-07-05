import io
import re

from docx import Document
from pypdf import PdfReader


class ParseError(Exception):
    pass


def extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(f"Could not read PDF: {exc}") from exc

    chunks: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(text.strip())
    combined = "\n".join(chunks).strip()
    if not combined:
        raise ParseError("PDF contained no extractable text.")
    return combined


def extract_docx_text(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(f"Could not read Word document: {exc}") from exc

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    combined = "\n".join(chunks).strip()
    if not combined:
        raise ParseError("Word document contained no extractable text.")
    return combined


def extract_uploaded_text(data: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return extract_docx_text(data)
    if name.endswith(".pdf"):
        return extract_pdf_text(data)
    raise ParseError(f"Unsupported file type: {filename}. Upload PDF or DOCX, or paste text.")


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def parse_document_text(
    *,
    pasted: str,
    uploaded_bytes: bytes | None,
    uploaded_filename: str | None = None,
    label: str,
    allowed_extensions: tuple[str, ...] = (".pdf",),
) -> str:
    if uploaded_bytes:
        name = (uploaded_filename or "").lower()
        if not any(name.endswith(ext) for ext in allowed_extensions):
            allowed = ", ".join(ext.lstrip(".") for ext in allowed_extensions)
            raise ParseError(f"{label} upload must be {allowed}.")
        return extract_uploaded_text(uploaded_bytes, uploaded_filename or "")
    if pasted and pasted.strip():
        return pasted.strip()
    allowed = ", ".join(ext.lstrip(".") for ext in allowed_extensions)
    raise ParseError(f"{label} is required. Paste text or upload {allowed}.")


def parse_optional_document_text(
    *,
    pasted: str,
    uploaded_bytes: bytes | None,
    uploaded_filename: str | None = None,
    label: str,
    allowed_extensions: tuple[str, ...] = (".pdf", ".docx"),
) -> str:
    """Like parse_document_text, but returns empty string when nothing is provided."""
    if uploaded_bytes:
        name = (uploaded_filename or "").lower()
        if not any(name.endswith(ext) for ext in allowed_extensions):
            allowed = ", ".join(ext.lstrip(".") for ext in allowed_extensions)
            raise ParseError(f"{label} upload must be {allowed}.")
        return extract_uploaded_text(uploaded_bytes, uploaded_filename or "")
    if pasted and pasted.strip():
        return pasted.strip()
    return ""
