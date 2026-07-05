import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from .schemas import CustomizedResumeContent, ResumeJobEntry, ResumeSectionContent

ACCENT_BLUE = RGBColor(0x1F, 0x5C, 0x99)
MUTED_GRAY = RGBColor(0x44, 0x44, 0x44)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT_NAME = "Arial"
CONTENT_WIDTH_IN = 7.0  # 8.5" page minus 0.75" left/right margins
MIN_RESUME_WORDS = 1100

_SOCIAL_LINK_FIELDS = (
    ("linkedin", "LinkedIn"),
    ("portfolio", "Portfolio"),
    ("kaggle", "Kaggle"),
    ("github", "GitHub"),
)


def _slug(text: str) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", text or "").strip()
    cleaned = re.sub(r"\s+", "-", cleaned)
    return cleaned or "Company"


def resume_download_filename(content: CustomizedResumeContent, profile: dict) -> str:
    first = profile.get("first_name") or "Alex"
    last = profile.get("last_name") or "Chen"
    company = _slug(content.company_name)
    return f"{first}-{last}-Resume-{company}.docx"


def _set_run_font(
    run,
    *,
    size_pt: float = 11,
    bold: bool = False,
    color: RGBColor = BLACK,
) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.append(r_fonts)


def _set_paragraph_bottom_border(paragraph, color_hex: str = "1F5C99", size_eighths: int = 6) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def count_resume_words(content: CustomizedResumeContent) -> int:
    parts: list[str] = []
    for section in content.sections:
        parts.extend(section.paragraphs)
        parts.extend(section.bullet_points)
        for entry in section.job_entries:
            parts.append(entry.title_line)
            parts.append(entry.date_range)
            parts.append(entry.location)
            parts.extend(entry.bullet_points)
    return len(re.findall(r"\b[\w'-]+\b", " ".join(parts)))


def _profile_social_links(profile: dict) -> list[tuple[str, str]]:
    links: list[tuple[str, str]] = []
    for field, label in _SOCIAL_LINK_FIELDS:
        url = str(profile.get(field) or "").strip()
        if url:
            links.append((label, url))
    return links


def _profile_certification_links(profile: dict) -> list[tuple[str, str]]:
    links: list[tuple[str, str]] = []
    for cert in profile.get("certifications") or []:
        if not isinstance(cert, dict):
            continue
        name = str(cert.get("name") or "").strip()
        url = str(cert.get("url") or "").strip()
        if name and url:
            links.append((name, url))
    return links


def _add_hyperlink(paragraph, text: str, url: str, *, size_pt: float = 11) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1F5C99")
    r_pr.append(color_el)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size_pt * 2)))
    r_pr.append(size_el)

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.append(r_fonts)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_centered_text_line(
    doc: Document,
    text: str,
    *,
    size_pt: float = 11,
    bold: bool = False,
    color: RGBColor = ACCENT_BLUE,
) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _set_run_font(run, size_pt=size_pt, bold=bold, color=color)


def _add_centered_link_row(doc: Document, links: list[tuple[str, str]]) -> None:
    if not links:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (label, url) in enumerate(links):
        if index:
            sep = para.add_run(" | ")
            _set_run_font(sep, size_pt=11, color=ACCENT_BLUE)
        _add_hyperlink(para, label, url)


def _add_header_from_profile(doc: Document, profile: dict) -> None:
    first = str(profile.get("first_name") or "").strip()
    last = str(profile.get("last_name") or "").strip()
    name = f"{first} {last}".strip()
    if name:
        _add_centered_text_line(doc, name, size_pt=16, bold=True, color=ACCENT_BLUE)

    location = str(profile.get("location") or "").strip()
    if location:
        _add_centered_text_line(doc, location, color=MUTED_GRAY)

    email = str(profile.get("email") or "").strip()
    phone = str(profile.get("phone") or "").strip()
    if email or phone:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if email:
            _add_hyperlink(para, email, f"mailto:{email}")
        if email and phone:
            sep = para.add_run(" | ")
            _set_run_font(sep, size_pt=11, color=ACCENT_BLUE)
        if phone:
            phone_run = para.add_run(phone)
            _set_run_font(phone_run, size_pt=11, color=ACCENT_BLUE)

    _add_centered_link_row(doc, _profile_social_links(profile))
    _add_centered_link_row(doc, _profile_certification_links(profile))


def _add_section_heading(doc: Document, heading: str) -> None:
    text = heading.strip().upper()
    if not text:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Twips(140)
    para.paragraph_format.space_after = Twips(20)
    run = para.add_run(text)
    _set_run_font(run, size_pt=12, bold=True, color=ACCENT_BLUE)
    _set_paragraph_bottom_border(para)


def _add_body_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Twips(20)
    para.paragraph_format.space_after = Twips(20)
    run = para.add_run(stripped)
    _set_run_font(run, size_pt=11)


def _add_bullet_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        return
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Twips(20)
    para.paragraph_format.space_after = Twips(20)
    _write_inline_bold(para, stripped)


def _write_inline_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size_pt=11, bold=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size_pt=11)


def _add_job_entry(doc: Document, entry: ResumeJobEntry) -> None:
    title = entry.title_line.strip()
    dates = entry.date_range.strip()
    if title:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Twips(20)
        para.paragraph_format.space_after = Twips(20)
        para.paragraph_format.tab_stops.add_tab_stop(
            Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
        )
        title_run = para.add_run(title)
        _set_run_font(title_run, size_pt=11, bold=True)
        if dates:
            para.add_run("\t")
            date_run = para.add_run(dates)
            _set_run_font(date_run, size_pt=11)

    location = entry.location.strip()
    if location:
        loc_para = doc.add_paragraph()
        loc_para.paragraph_format.space_after = Twips(10)
        loc_run = loc_para.add_run(location)
        _set_run_font(loc_run, size_pt=11, color=MUTED_GRAY)

    for bullet in entry.bullet_points:
        _add_bullet_paragraph(doc, bullet)


def _add_section(doc: Document, section: ResumeSectionContent) -> None:
    _add_section_heading(doc, section.heading)
    for paragraph in section.paragraphs:
        _add_body_paragraph(doc, paragraph)
    for entry in section.job_entries:
        _add_job_entry(doc, entry)
    for bullet in section.bullet_points:
        _add_bullet_paragraph(doc, bullet)


def _configure_page(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)

    page = doc.sections[0]
    page.page_width = Inches(8.5)
    page.page_height = Inches(11)
    page.left_margin = Inches(0.75)
    page.right_margin = Inches(0.75)
    page.top_margin = Inches(0.625)
    page.bottom_margin = Inches(0.625)


def build_resume_docx(content: CustomizedResumeContent, profile: dict | None = None) -> bytes:
    doc = Document()
    _configure_page(doc)

    if profile:
        _add_header_from_profile(doc, profile)
        if content.sections:
            doc.add_paragraph("")
    elif content.header_lines:
        for index, line in enumerate(content.header_lines):
            text = line.strip()
            if not text:
                continue
            _add_centered_text_line(
                doc,
                text,
                size_pt=16 if index == 0 else 11,
                bold=index == 0,
                color=MUTED_GRAY if index == 1 else ACCENT_BLUE,
            )
        if content.sections:
            doc.add_paragraph("")

    for section in content.sections:
        _add_section(doc, section)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
